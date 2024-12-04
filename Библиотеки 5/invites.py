from docx import Document
from docx.shared import Inches
from sys import stdin

data = [input('Место: '), input('Время: '), []]
[data[2].append(' '.join(person.split())) for person in stdin]

document = Document()
for invite in data[2]:
    document.add_heading('Приглашение', 0)
    p = document.add_paragraph('Приветствую вас, ')
    p.add_run(invite+'. ').bold = True
    p.add_run('Рад пригласить вас на очень значимое для меня мероприятие, которое пройдёт в ')
    p.add_run(data[0]+' в '+data[1]+'.').bold = True
    p = document.add_paragraph('С наилучшими пожеланиями, ваш друг.')
    document.add_page_break()

document.save('test.docx')